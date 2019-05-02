import docx
from docx.shared import Inches
import sys, csv

csvInput = 'input.csv'

VERSION = '0.1'
print('prop65-word.py: Version ' + VERSION)

# INSTRUCTIONS: input.csv should be in the same directory as this script.

# Do not include a header row in input.csv
# Column A of input.csv should have file names (do not include .docx)
# Column B of input.csv should contain Prop 65 warning messages to be printed on documents corresponding to that file name.
# input.csv should NOT be in UTF-8 format

# Example:

# [    A     ][                           B                      ]
# [ ABC12345 ][        Warning: This is the Prop 65 Warning      ]
# [ BCD23456 ][ Warning: This is another Prop 65 Warning Message ]

# Output:

# ABC12345.docx: /!\ Warning: This is the Prop 65 Warning
# BCD23456.docx: /!\ Warning: This is another Prop 65 Warning Message

#  To bulk create PDFs after .doc files are generated, use PDFCreator/PDFArchitect and set as Default Printer
#  Make sure settings for PDFCreator/PDFArchitect are set to automate the file saving process (no dialog box)
#  Then, highlight multiple .doc files, right click, choose Print, and PDFs should be generated

try:
    csvInputFileObj = open(csvInput, 'r')
    csvReader = csv.reader(csvInputFileObj)
    
except FileNotFoundError:
    print('ERROR: Could not find CSV input file named "input.csv" in same directory as this script.')
    sys.exit()


for row in csvReader:
    
    filename = row[0]
    text = row[1]
    
    d = docx.Document()
    p = d.add_paragraph()
    r = p.add_run()

    r.add_picture('6pt.png',width=Inches(0.25))
    r.add_text('  ')
    r.add_text(text)

    d.save(filename + '.docx')
    print(filename + '.docx saved')
    
csvInputFileObj.close()  
